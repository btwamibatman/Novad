from __future__ import annotations

from collections.abc import Callable
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
from uuid import uuid4

import pymupdf
from fastapi import UploadFile

from app.core.config import settings
from app.services.text_analysis import extract_pdf_pages_with_ocr

ProgressCallback = Callable[[int, str], None]


class DocumentToolError(RuntimeError):
    pass


async def save_word_upload(file: UploadFile) -> tuple[str, str, str]:
    filename = Path((file.filename or "").replace("\\", "/")).name.strip()
    extension = Path(filename).suffix.lower()
    if not filename or extension not in {".docx", ".doc", ".odt"}:
        raise DocumentToolError("Choose a DOCX, DOC or ODT document")
    directory = Path(settings.storage_dir).parent / "tools" / "inputs"
    directory.mkdir(parents=True, exist_ok=True)
    path = directory / f"{uuid4()}{extension}"
    size = 0
    try:
        with path.open("wb") as output:
            while chunk := await file.read(1024 * 1024):
                size += len(chunk)
                if size > settings.max_upload_size_bytes:
                    raise DocumentToolError("File is too large")
                output.write(chunk)
        if size == 0:
            raise DocumentToolError("Uploaded file is empty")
    except Exception:
        path.unlink(missing_ok=True)
        raise
    content_type = {
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".odt": "application/vnd.oasis.opendocument.text",
    }[extension]
    return filename, content_type, str(path.resolve())


def compress_pdf(
    source: Path,
    destination: Path,
    mode: str,
    progress: ProgressCallback | None = None,
) -> dict:
    if mode not in {"low", "recommended", "extreme"}:
        raise DocumentToolError("Unknown compression mode")
    destination.parent.mkdir(parents=True, exist_ok=True)
    _progress(progress, 10, "preparing")
    if mode == "low":
        _lossless_rewrite(source, destination)
        pipeline = "pymupdf-lossless"
    else:
        _ghostscript_compress(source, destination, mode)
        pipeline = "ghostscript-balanced" if mode == "recommended" else "ghostscript-small"
    _progress(progress, 85, "validating")
    _validate_pdf(destination)

    original_size = source.stat().st_size
    result_size = destination.stat().st_size
    if result_size >= original_size:
        destination.unlink(missing_ok=True)
        shutil.copyfile(source, destination)
        result_size = original_size
        pipeline += "-original-retained"
    savings = 0 if not original_size else round((1 - result_size / original_size) * 100, 1)
    _progress(progress, 100, "completed")
    return {
        "mode": mode,
        "pipeline": pipeline,
        "original_size_bytes": original_size,
        "result_size_bytes": result_size,
        "savings_percent": max(savings, 0),
    }


def word_to_pdf(
    source: Path,
    destination: Path,
    progress: ProgressCallback | None = None,
) -> dict:
    executable = _find_executable("soffice", "libreoffice")
    if executable is None:
        raise DocumentToolError(
            "LibreOffice is unavailable. Run the service through the Docker image."
        )
    destination.parent.mkdir(parents=True, exist_ok=True)
    _progress(progress, 15, "preparing")
    with tempfile.TemporaryDirectory(prefix="document-convert-") as temp:
        output_dir = Path(temp) / "output"
        profile_dir = Path(temp) / "profile"
        output_dir.mkdir()
        command = [
            executable,
            "--headless",
            f"-env:UserInstallation={profile_dir.resolve().as_uri()}",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            str(output_dir),
            str(source),
        ]
        _progress(progress, 35, "converting")
        _run(command, timeout=300)
        candidates = list(output_dir.glob("*.pdf"))
        if len(candidates) != 1:
            raise DocumentToolError("LibreOffice did not produce a PDF file")
        shutil.move(str(candidates[0]), destination)
    _progress(progress, 90, "validating")
    _validate_pdf(destination)
    _progress(progress, 100, "completed")
    return {"pipeline": "libreoffice-local", "editable_source": True}


def pdf_to_word(
    source: Path,
    destination: Path,
    progress: ProgressCallback | None = None,
) -> dict:
    try:
        from docx import Document as WordDocument
        from docx.enum.text import WD_BREAK
    except ImportError as error:
        raise DocumentToolError("python-docx is not installed") from error

    _progress(progress, 5, "extracting")

    def extraction_progress(completed: int, total: int, stage: str) -> None:
        percent = 10 + int((completed / max(total, 1)) * 65)
        _progress(progress, percent, "ocr" if stage == "ocr" else "extracting")

    pages = extract_pdf_pages_with_ocr(source, progress_callback=extraction_progress)
    if not pages or not any(page.text.strip() for page in pages):
        raise DocumentToolError("No readable text was found in the PDF")
    _progress(progress, 80, "building_word")
    document = WordDocument()
    for index, page in enumerate(pages):
        if index:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        for block in re.split(r"\n\s*\n", page.text.strip()):
            if block.strip():
                document.add_paragraph(block.strip())
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)
    ocr_pages = sum(page.extraction_method.startswith("ocr") for page in pages)
    _progress(progress, 100, "completed")
    return {
        "pipeline": "local-extraction-python-docx",
        "page_count": len(pages),
        "ocr_page_count": ocr_pages,
        "beta": True,
    }


def render_pdf_page(source: Path, page_number: int, dpi: int = 120) -> bytes:
    try:
        with pymupdf.open(source) as document:
            if page_number < 1 or page_number > document.page_count:
                raise DocumentToolError("Page does not exist")
            page = document.load_page(page_number - 1)
            pixmap = page.get_pixmap(matrix=pymupdf.Matrix(dpi / 72, dpi / 72), alpha=False)
            return pixmap.tobytes("png")
    except DocumentToolError:
        raise
    except Exception as error:
        raise DocumentToolError("Unable to render PDF page") from error


def _lossless_rewrite(source: Path, destination: Path) -> None:
    try:
        with pymupdf.open(source) as document:
            if document.needs_pass:
                raise DocumentToolError("Password-protected PDFs are not supported")
            document.save(
                destination,
                garbage=4,
                clean=True,
                deflate=True,
                deflate_images=True,
                deflate_fonts=True,
                use_objstms=1,
            )
    except DocumentToolError:
        raise
    except Exception as error:
        raise DocumentToolError("PDF optimization failed") from error


def _ghostscript_compress(source: Path, destination: Path, mode: str) -> None:
    executable = _find_executable("gswin64c", "gswin32c", "gs")
    if executable is None:
        raise DocumentToolError(
            "Ghostscript is unavailable. Run the service through the Docker image."
        )
    if mode == "recommended":
        resolution, mono_resolution, jpeg_quality = 150, 300, 80
    else:
        resolution, mono_resolution, jpeg_quality = 96, 200, 55
    command = [
        executable,
        "-dSAFER",
        "-dBATCH",
        "-dNOPAUSE",
        "-sDEVICE=pdfwrite",
        "-dCompatibilityLevel=1.7",
        "-dDetectDuplicateImages=true",
        "-dCompressFonts=true",
        "-dSubsetFonts=true",
        "-dDownsampleColorImages=true",
        "-dDownsampleGrayImages=true",
        "-dDownsampleMonoImages=true",
        "-dColorImageDownsampleType=/Bicubic",
        "-dGrayImageDownsampleType=/Bicubic",
        f"-dColorImageResolution={resolution}",
        f"-dGrayImageResolution={resolution}",
        f"-dMonoImageResolution={mono_resolution}",
        f"-dJPEGQ={jpeg_quality}",
        f"-sOutputFile={destination}",
        str(source),
    ]
    _run(command, timeout=300)


def _validate_pdf(path: Path) -> None:
    try:
        with pymupdf.open(path) as document:
            if document.page_count < 1:
                raise DocumentToolError("Resulting PDF has no pages")
    except DocumentToolError:
        raise
    except Exception as error:
        raise DocumentToolError("Resulting PDF is invalid") from error


def _find_executable(*names: str) -> str | None:
    for name in names:
        path = shutil.which(name)
        if path:
            return path
    return None


def _run(command: list[str], timeout: int) -> None:
    try:
        result = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=timeout,
            check=False,
        )
    except (OSError, subprocess.TimeoutExpired) as error:
        raise DocumentToolError("Local document processing timed out or could not start") from error
    if result.returncode != 0:
        detail = (result.stderr or result.stdout).strip()[-800:]
        raise DocumentToolError(detail or "Local document processing failed")


def _progress(callback: ProgressCallback | None, percent: int, stage: str) -> None:
    if callback:
        callback(max(0, min(percent, 100)), stage)
